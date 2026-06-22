import re
import os
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def set_run_font(run, font_name, size_pt, bold=False, italic=False):
	"""设置Run的字体和字号，同时适配中西文字体，兼容国家公文标准"""
	run.font.size = Pt(size_pt)
	run.bold = bold
	run.italic = italic
	# 西文字体（英文、数字等）设为 Times New Roman
	run.font.name = 'Times New Roman'
	# 关键：必须显式设置东亚中文字体属性
	rPr = run._element.get_or_add_rPr()
	rFonts = rPr.get_or_add_rFonts()
	rFonts.set(qn('w:eastAsia'), font_name)
	# 同时设置西文/ASCII 字体属性，确保英文和数字用 Times New Roman
	rFonts.set(qn('w:ascii'), 'Times New Roman')
	rFonts.set(qn('w:hAnsi'), 'Times New Roman')
	rFonts.set(qn('w:cs'), 'Times New Roman')

def apply_paragraph_format(p, alignment, first_line_indent_pt=0, line_spacing_pt=29, space_before=0, space_after=0):
	"""应用标准的公文段落格式，保证每页22行"""
	p.alignment = alignment
	p_format = p.paragraph_format
	p_format.space_before = Pt(space_before)
	p_format.space_after = Pt(space_after)
	# 设置行距为固定值（公文标准每页22行，对应约29pt固定行距）
	p_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
	p_format.line_spacing = Pt(line_spacing_pt)
	if first_line_indent_pt > 0:
		p_format.first_line_indent = Pt(first_line_indent_pt)

def parse_markdown(md_text):
	"""解析Markdown文本，提取Front Matter元数据和公文正文结构"""
	lines = md_text.splitlines()
			
	metadata = {'落款': '', '日期': ''}
	body_lines = []
	in_front_matter = False
	front_matter_checked = False
			
	for line in lines:
		stripped = line.strip()
		# 检测 Front Matter 边界 ---
		if stripped == '---':
			if not front_matter_checked and not in_front_matter:
				in_front_matter = True
				continue
			elif in_front_matter:
				in_front_matter = False
				front_matter_checked = True
				continue
		
		if in_front_matter:
			# 解析 YAML 键值对
			if ':' in line:
				key, val = line.split(':', 1)
				metadata[key.strip()] = val.strip()
			elif '：' in line:
				key, val = line.split('：', 1)
				metadata[key.strip()] = val.strip()
		else:
			body_lines.append(line)
			
	# 将连续的非空行组合成自然段落
	paragraphs = []
	current_para = []
			
	for line in body_lines:
		if line.strip() == '':
			if current_para:
				paragraphs.append('\n'.join(current_para))
				current_para = []
		else:
			current_para.append(line)
	if current_para:
		paragraphs.append('\n'.join(current_para))
		
	# 精准识别各段落开头的标识符，映射到对应的公文层级
	elements = []
	for para in paragraphs:
		text = para.strip()
		if text.startswith('# '):
			# Markdown # 映射为公文大标题
			elements.append(('TITLE', text[2:].strip()))
		elif re.match(r'^[一二三四五六七八九十百]+、', text):
			elements.append(('H1', text))
		elif re.match(r'^[（(][一二三四五六七八九十百]+[）)]', text):
			elements.append(('H2', text))
		elif re.match(r'^\d+[\.、]', text):
			elements.append(('H3', text))
		elif re.match(r'^[（(]\d+[）)]', text):
			elements.append(('H4', text))
		else:
			elements.append(('BODY', text))
			
	return metadata, elements

def convert_md_to_gov_docx(md_content, output_path):
	"""将解析后的Markdown内容渲染成符合国家公文标准的 Word docx 文档"""
	metadata, elements = parse_markdown(md_content)
			
	doc = Document()
			
	# 1. 设置标准的公文纸张大小与页边距 (GB/T 9704-2012)
	section = doc.sections[0]
	section.page_width = Mm(210)   # A4 宽
	section.page_height = Mm(297)  # A4 高
	section.top_margin = Mm(37)	# 天头 37mm
	section.bottom_margin = Mm(35) # 地脚 35mm
	section.left_margin = Mm(28)   # 订口 28mm
	section.right_margin = Mm(26)  # 切口 26mm
			
	# 2. 写入公文要素
	for type_, text in elements:
		p = doc.add_paragraph()
		
		if type_ == 'TITLE':
			# 公文标题：2号小标宋体，居中，加大字体的行间距
			apply_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, line_spacing_pt=35, space_before=12, space_after=18)
			run = p.add_run(text)
			set_run_font(run, '方正小标宋简体', 22, bold=True)
			
		elif type_ == 'H1':
			# 一级标题：3号黑体，左空二字（首行缩进32pt）
			apply_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_pt=32, space_before=4, space_after=4)
			run = p.add_run(text)
			set_run_font(run, '黑体', 16, bold=True)
			
		elif type_ == 'H2':
			# 二级标题：3号楷体，左空二字
			apply_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_pt=32, space_before=2, space_after=2)
			run = p.add_run(text)
			set_run_font(run, '楷体', 16)
			
		elif type_ == 'H3':
			# 三级标题：3号仿宋加粗，左空二字
			apply_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_pt=32, space_before=1, space_after=1)
			run = p.add_run(text)
			set_run_font(run, '仿宋', 16, bold=True)
			
		elif type_ == 'H4':
			# 四级标题：3号仿宋，左空二字
			apply_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_pt=32)
			run = p.add_run(text)
			set_run_font(run, '仿宋', 16)
			
		elif type_ == 'BODY':
			# 普通正文：3号仿宋，左空二字，行距固定29pt
			apply_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_pt=32)
			run = p.add_run(text)
			set_run_font(run, '仿宋', 16)
			
	# 3. 添加尾部发文机关署名（落款）和成文日期
	if metadata['落款'] or metadata['日期']:
		# 按标准留出空行隔开正文
		p_space = doc.add_paragraph()
		apply_paragraph_format(p_space, WD_ALIGN_PARAGRAPH.LEFT)
		
		if metadata['落款']:
			p_org = doc.add_paragraph()
			apply_paragraph_format(p_org, WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)
			run_org = p_org.add_run(metadata['落款'])
			set_run_font(run_org, '仿宋', 16)
			
		if metadata['日期']:
			p_date = doc.add_paragraph()
			apply_paragraph_format(p_date, WD_ALIGN_PARAGRAPH.RIGHT)
			run_date = p_date.add_run(metadata['日期'])
			set_run_font(run_date, '仿宋', 16)
			
	doc.save(output_path)
	print(f"成功编译公文并保存至: {output_path}")


if __name__ == '__main__':
	import sys

	if len(sys.argv) < 2:
		print("用法: python ofdoc.py <输入文件.md>")
		sys.exit(1)

	input_path = sys.argv[1]
	if not os.path.isfile(input_path):
		print(f"错误: 文件不存在 - {input_path}")
		sys.exit(1)

	with open(input_path, 'r', encoding='utf-8') as f:
		md_content = f.read()

	# 将输入文件名后缀替换为 .docx 作为输出文件名
	output_path = os.path.splitext(input_path)[0] + '.docx'
	convert_md_to_gov_docx(md_content, output_path)